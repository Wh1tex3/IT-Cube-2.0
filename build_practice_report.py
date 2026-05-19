# -*- coding: utf-8 -*-
"""Сборка отчёта ППС: правки шаблона + вставка текста по проекту IT Cube 2.0."""
from __future__ import annotations

import os
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph

SRC = os.path.join(r"c:\Users\Данил\Desktop", "Тит_содерж_введ_закл_ППС.docx")
DST = os.path.join(r"c:\Users\Данил\Desktop", "Отчёт_ППС_Шищова_ГБУ_Приоритет.docx")

ORG = "ГБУ ДО ЛО «Центр образования «Приоритет»»"
OLD_ORG = "ООО «Дружба»"


def set_run_font(run, size_pt=14, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")


def format_body_paragraph(p):
    pf = p.paragraph_format
    pf.first_line_indent = Cm(1.25)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    for r in p.runs:
        set_run_font(r, 14, False)


def add_heading_before(body_ref: Paragraph, text: str):
    p = OxmlElement("w:p")
    body_ref._element.addprevious(p)
    new_p = Paragraph(p, body_ref._parent)
    run = new_p.add_run(text)
    set_run_font(run, 14, bold=True)
    new_p.paragraph_format.space_before = Pt(12)
    new_p.paragraph_format.space_after = Pt(6)
    new_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    new_p.paragraph_format.first_line_indent = Cm(0)
    new_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    new_p.paragraph_format.line_spacing = 1.5
    return new_p


def add_body_before(body_ref: Paragraph, text: str):
    p = OxmlElement("w:p")
    body_ref._element.addprevious(p)
    new_p = Paragraph(p, body_ref._parent)
    new_p.add_run(text)
    format_body_paragraph(new_p)
    return new_p


def replace_in_cell(cell, old: str, new: str):
    for p in cell.paragraphs:
        if old in p.text:
            full = p.text.replace(old, new)
            for r in p.runs:
                r.text = ""
            if p.runs:
                p.runs[0].text = full
            else:
                p.add_run(full)
            for r in p.runs:
                set_run_font(r, 14, False)


def patch_tables(doc: Document):
    t0 = doc.tables[0].rows[0].cells
    replace_in_cell(t0[1], "Иванов И.И.", "Шищова Е.Д.")
    for p in t0[1].paragraphs:
        for r in p.runs:
            set_run_font(r, 14, False)

    org_cell = doc.tables[3].rows[1].cells[0]
    replace_in_cell(org_cell, OLD_ORG, ORG)
    for p in org_cell.paragraphs:
        for r in p.runs:
            set_run_font(r, 14, False)

    toc = doc.tables[6]
    fills = {
        4: "1.2 Функциональные возможности веб-приложения",
        7: "2.2 Структура клиентского приложения и разделы интерфейса",
        8: "2.3 Модель данных: пользователи, объединения, инструкции и коллекции",
        9: "2.4 Регистрация, вход в систему и разграничение ролей",
        10: "2.5 Раздел «Инструкции»: фильтрация, карточки, модальное окно и завершение задания",
        11: "2.6 Раздел «Статистика», расчёт уровня и опыта, таблица лидеров",
        12: "2.7 Панель администратора: инструкции, коллекции, пользователи",
        13: "2.8 Сохранение состояния в localStorage и загрузка данных из каталога Projectlego",
        16: "3.1 Общие сведения об информационной безопасности",
        17: "3.2 Информационная безопасность в разрабатываемом веб-приложении",
    }
    for ri, text in fills.items():
        cell = toc.rows[ri].cells[0]
        cell.text = text
        for p in cell.paragraphs:
            for r in p.runs:
                set_run_font(r, 14, False)


def replace_in_paragraphs(doc: Document, old: str, new: str):
    for p in doc.paragraphs:
        if old in p.text:
            txt = p.text.replace(old, new)
            for r in p.runs:
                r.text = ""
            if p.runs:
                p.runs[0].text = txt
            else:
                p.add_run(txt)
            for r in p.runs:
                set_run_font(r, 14, False)


def insert_main_body_before_conclusion(doc: Document):
    target = None
    for p in doc.paragraphs:
        if p.text.strip() == "Заключение":
            target = p
            break
    if target is None:
        raise RuntimeError("Не найден абзац «Заключение»")

    blocks: list[tuple[str, str]] = []

    def H(s: str):
        blocks.append(("h", s))

    def B(s: str):
        blocks.append(("b", s))

    H("1 Теоретическая часть")
    H("1.1 Описание объекта проектирования")
    B(
        "Объектом проектирования является одностраничное веб-приложение "
        "«Сайт для роботов — образовательная платформа», реализованное в виде HTML-страницы "
        "с подключением JavaScript-фреймворка Vue 3 и библиотеки стилей Bootstrap 5. "
        "Приложение предназначено для организации учебного процесса по курсу робототехники "
        "на базе конструктора LEGO WeDo 2.0: просмотр инструкций по сборке, учёт выполненных работ, "
        "отображение статистики и рейтинга учащихся, а также ведение учётных записей преподавателя "
        "и администратора. Серверная часть и база данных в проекте не используются: данные "
        "хранятся в браузере пользователя (localStorage)."
    )
    H("1.2 Функциональные возможности веб-приложения")
    B(
        "В приложении реализованы следующие возможности, непосредственно отражённые в исходном коде. "
        "Экран входа и регистрации: при регистрации выбирается роль «Учащийся» или «Педагог»; при входе дополнительно "
        "доступна роль «Модератор». Учащийся вводит код объединения; педагог задаёт название нового объединения. "
        "Раздел «Инструкции»: список инструкций и коллекций, фильтры по категории, сложности, коллекции и текстовому поиску, "
        "карточки инструкций с изображением, открытие подробностей в модальном окне Bootstrap, завершение инструкции "
        "с начислением опыта и вводом кода подтверждения от преподавателя для учащегося. Раздел «Статистика»: отображение опыта, "
        "уровня, прогресса, таблица лидеров с переключением младшей и старшей лиги. Раздел «Профиль»: "
        "данные пользователя, выбор пола для аватара, сводная статистика. Для ролей администратора и модератора "
        "доступна «Панель администратора» с вкладками для создания инструкций и коллекций, "
        "генерации кода подтверждения, управления пользователями (в том числе срок действия учётной записи и деактивация). "
        "Реализовано переключение светлой и тёмной темы с сохранением выбора в localStorage. "
        "При запуске выполняется попытка импорта описаний наборов из локального каталога Projectlego "
        "путём загрузки и разбора HTML-страниц."
    )
    H("2 Проектная часть")
    H("2.1 Обоснование выбора средств проектирования")
    B(
        "Для реализации клиентской части выбран Vue 3 в варианте подключения через CDN без сборщика: "
        "это сокращает объём настройки окружения при учебной разработке и достаточно для одностраничного "
        "приложения с реактивным состоянием. В качестве основы макета и компонентов интерфейса использован "
        "Bootstrap 5 (сетка, формы, навигационная панель, модальные окна). Для хранения данных без сервера "
        "используется Web Storage API (localStorage), что соответствует выбранной архитектуре прототипа. "
        "Стилизация выполнена отдельным файлом styles.css поверх стандартных классов Bootstrap."
    )
    H("2.2 Структура клиентского приложения и разделы интерфейса")
    B(
        "Исходные файлы приложения: index.html (разметка и привязка директив Vue), app.js (объект состояния, "
        "вычисляемые свойства, методы обработки действий пользователя, загрузка и сохранение данных), "
        "styles.css (оформление экранов). Переключение экранов выполняется по полю currentSection "
        "(значения auth, instructions, statistics, profile, admin); навигационная панель скрывается на экране авторизации."
    )
    H("2.3 Модель данных: пользователи, объединения, инструкции и коллекции")
    B(
        "В состоянии приложения хранятся массивы users, groups, instructions и collections. "
        "Пользователь содержит идентификатор, ФИО, логин и пароль, роль (user, admin, moderator), "
        "идентификатор объединения, опыт, списки завершённых инструкций и результатов по опыту, "
        "признак активности и срок действия учётной записи, пол и служебные поля для кода подтверждения. "
        "Объединение задаёт код для вступления учащихся, наименование и ФИО педагога по умолчанию. "
        "Инструкция включает заголовок, категории, привязку к коллекции и группе, параметры сложности "
        "(число слайдов, сложные соединения, сложность программы, бонусы за самостоятельную сборку, "
        "программирование и устранение неисправностей), уровень сложности easy, medium или hard, адрес изображения, "
        "признаки мотора и датчиков, формат материала. Коллекция содержит имя и идентификатор группы."
    )
    H("2.4 Регистрация, вход в систему и разграничение ролей")
    B(
        "Вход выполняется методом handleLogin: проверяется совпадение логина, пароля и роли с записью "
        "в массиве users и признак active. Регистрация handleRegister: для педагога создаётся новая группа "
        "со случайным четырёхзначным кодом; для учащегося проверяется код существующего объединения. "
        "Проверяется уникальность логина. После успешной регистрации или входа текущий пользователь копируется "
        "в currentUser и открывается раздел instructions. Вычисляемое свойство isAdminLike определяет "
        "отображение пункта «Панель администратора» для ролей admin и moderator."
    )
    H("2.5 Раздел «Инструкции»: фильтрация, карточки, модальное окно и завершение задания")
    B(
        "Список инструкций группы формируется вычисляемым свойством groupInstructions с учётом фильтров "
        "filters (категория, сложность, коллекция, только невыполненные, поиск). Карточка открывает модальное окно "
        "instructionModal. Метод computeMaxExp рассчитывает максимальный опыт по параметрам инструкции. "
        "Для учащегося завершение с кодом completeInstructionWithCode сравнивает введённый код с teacherConfirmCode "
        "активных администраторов группы; при совпадении вызывается completeInstruction с записью опыта и "
        "идентификатора инструкции в completedInstructions. Для администратора предусмотрено завершение без кода "
        "с выбором начисляемого опыта до максимума."
    )
    H("2.6 Раздел «Статистика», расчёт уровня и опыта, таблица лидеров")
    B(
        "Уровень пользователя вычисляется как единица плюс целая часть от деления числа завершённых инструкций на пять. "
        "Прогресс до следующего уровня отображается через остаток от деления на пять. В таблице лидеров используется "
        "переключатель leaderboardLeague (junior или senior) и отсортированные списки учащихся группы; для позиции "
        "в рейтинге рассчитывается число инструкций до обгона предыдущего участника. Отображаются также вычисляемые "
        "проценты для индикаторов прогресса уровня и позиции в таблице."
    )
    H("2.7 Панель администратора: инструкции, коллекции, пользователи")
    B(
        "Панель переключается полем adminTab. Создание инструкции createInstruction формирует объект с полями формы "
        "instructionForm и сохраняет его в массив instructions. Создание коллекции createCollection добавляет запись "
        "в collections. Для пользователей доступны продление срока activeUntil, установка даты, деактивация deleteUser, "
        "генерация кода подтверждения generateTeacherConfirmCode. Изменения фиксируются вызовом saveState."
    )
    H("2.8 Сохранение состояния в localStorage и загрузка данных из каталога Projectlego")
    B(
        "Метод saveState сериализует users, instructions, collections и groups в JSON и записывает в ключ "
        "localStorage «robot-site-state». Метод loadState читает и восстанавливает состояние, при ошибке или "
        "отсутствии данных вызывается seedInitialData с демонстрационными пользователями. "
        "Метод importProjectlego загружает HTML-файлы из каталога Projectlego, разбирает DOM и добавляет инструкции "
        "и коллекции при отсутствии дубликатов; флаг pl-imported-v1 предотвращает повторный импорт."
    )
    H("2.9 Разработка интерфейса пользователя")
    B(
        "Интерфейс построен на компонентах Bootstrap (navbar, container, cards, modal) с кастомными классами в styles.css: "
        "оформление экрана авторизации с фоновым изображением и сеткой auth-layout, стили карточек инструкций и коллекций, "
        "кнопок, таблицы лидеров и административных форм. Подключён шрифт Nunito с Google Fonts. "
        "Значок робота в шапке реализован символом emoji в круге."
    )
    H("3 Информационная безопасность")
    H("3.1 Общие сведения об информационной безопасности")
    B(
        "Под информационной безопасностью понимают защищённость информации и поддерживающей её инфраструктуры "
        "от случайных или преднамеренных воздействий, нарушающих конфиденциальность, целостность и доступность данных. "
        "Она необходима для снижения рисков утечки персональных и служебных сведений, искажения или уничтожения информации, "
        "а также для обеспечения непрерывности работы информационных систем. На организационном уровне вводят политики "
        "доступа, обучение пользователей, учёт носителей и инцидентов; на техническом — применяют разграничение доступа, "
        "шифрование каналов и хранилищ, резервное копирование, антивирусную защиту, обновления программного обеспечения и мониторинг."
    )
    H("3.2 Информационная безопасность в разрабатываемом веб-приложении")
    B(
        "В текущей реализации приложение не использует серверную аутентификацию: логины и пароли хранятся "
        "в открытом виде внутри JSON в localStorage браузера, что делает их доступными при физическом доступе "
        "к устройству или при выполнении стороннего сценария в контексте страницы. Проверка ролей и прав "
        "выполняется только на стороне клиента, поэтому изменение кода в инструменте разработчика теоретически "
        "позволяет обойти ограничения интерфейса. Положительным элементом является использование кода подтверждения "
        "преподавателя при зачёте инструкции учащимся и возможность деактивации учётной записи администратором. "
        "Для повышения безопасности в промышленной версии потребовались бы сервер, хэширование паролей, "
        "сеансовая аутентификация и передача данных по протоколу HTTPS."
    )

    ref = target
    for kind, text in reversed(blocks):
        if kind == "h":
            ref = add_heading_before(ref, text)
        else:
            ref = add_body_before(ref, text)


def append_bibliography(doc: Document):
    h = doc.add_paragraph()
    h.add_run("Список использованных источников")
    set_run_font(h.runs[0], 14, bold=True)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.first_line_indent = Cm(0)
    h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    h.paragraph_format.line_spacing = 1.5

    bib = (
        "1. Vue.js v3. Документация [Электронный ресурс]. — URL: https://vuejs.org/ (дата обращения: 12.05.2026).\n"
        "2. Bootstrap v5.3. Документация [Электронный ресурс]. — URL: https://getbootstrap.com/docs/5.3/ "
        "(дата обращения: 12.05.2026).\n"
        "3. MDN Web Docs. Window.localStorage [Электронный ресурс]. — URL: "
        "https://developer.mozilla.org/ru/docs/Web/API/Window/localStorage (дата обращения: 12.05.2026).\n"
        "4. HTML Living Standard. WHATWG [Электронный ресурс]. — URL: https://html.spec.whatwg.org/ "
        "(дата обращения: 12.05.2026)."
    )
    p = doc.add_paragraph()
    p.add_run(bib)
    format_body_paragraph(p)


def main():
    shutil.copy2(SRC, DST)
    doc = Document(DST)
    patch_tables(doc)
    replace_in_paragraphs(doc, OLD_ORG, ORG)
    insert_main_body_before_conclusion(doc)
    append_bibliography(doc)
    doc.save(DST)
    print("Saved:", DST)


if __name__ == "__main__":
    main()
